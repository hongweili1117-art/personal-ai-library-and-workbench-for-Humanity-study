"""
学术 PDF 阅读工作台 · 网页版 (Flask + PDF.js) —— 多 AI 接口增强版

本地运行（在 Anaconda Prompt 里）：
  pip install -r requirements.txt
  python app.py
然后浏览器会自动打开 http://127.0.0.1:5000

推荐 Windows 启动方式：双击「启动阅读工作台.vbs」或「launch_app.pyw」，不会弹出黑色命令行窗口。

离线使用 PDF.js（可选）：把 pdf.min.js 和 pdf.worker.min.js (v3.11.174)
放进 app.py 同目录下的 lib 文件夹，即可在断网时也能阅读。
"""

import json
import re
import threading
import webbrowser
import math
import uuid
import base64
import os
import sys
import traceback
import shutil
import urllib.request
import urllib.error
import html.parser as _html_parser
from collections import Counter
from datetime import datetime
from pathlib import Path

from flask import Flask, request, jsonify, send_file, Response

try:
    import pymupdf as fitz
except ImportError:
    import fitz

from config import *  # 路径 / AI 供应商与模型 / OCR 提示词 / 停用词 / IRON 等常量


from storage import *  # 数据存取层（library/highlights/notes/ai_collection/ocr 读写 + PDF 路径解析）


_text_cache = {}




def book_pages_text(book):
    key = book["id"]
    if key not in _text_cache:
        if (book.get("kind") or "pdf") in ("text", "snippet"):
            _text_cache[key] = _read_text_pages(book)
            return _text_cache[key]
        doc = fitz.open(resolve_pdf_path(book))
        pages = [doc[i].get_text() for i in range(doc.page_count)]
        doc.close()
        # 对扫描件或图片型 PDF：如果某页没有内嵌文字，但用户已经做过 OCR，
        # 就用 OCR 缓存补上，这样“原文”、RAG 和摘要都能继续使用。
        ocr = safe_ocr_cache()
        for i in range(len(pages)):
            if not (pages[i] or "").strip():
                cached = ocr.get(f"{key}:{i+1}")
                if isinstance(cached, dict) and (cached.get("text") or "").strip():
                    pages[i] = cached.get("text") or ""
        _text_cache[key] = pages
    return _text_cache[key]


def notes_list(alln, bid):
    raw = alln.get(bid, [])
    if isinstance(raw, str):  # 兼容旧版（单条字符串笔记）
        return [{"id": "n1", "title": "笔记", "body": raw}] if raw.strip() else []
    return raw


def clean_text(t):
    """缝合跨行连字符：Verfassungs-\nbeschwerde -> Verfassungsbeschwerde"""
    return re.sub(r'-\n\s*', '', t or '')


def with_glossary(sys, glossary):
    """把术语表拼到 system prompt 末尾，约束模型用法。"""
    g = (glossary or "").strip()
    if g:
        return (sys + "\n\n请严格遵守以下术语对照表（每行一条，左为原文术语、右为指定译法或说明）。"
                "出现这些术语时必须使用指定译法，不要换成同义词：\n" + g)
    return sys


def _tokenize_query(q):
    """轻量关键词切分：用于本地全文检索，不依赖额外向量库。"""
    toks = re.findall(r"[\w\u4e00-\u9fff]+", (q or "").lower(), flags=re.UNICODE)
    return [t for t in toks if len(t) > 1 and t not in STOPWORDS]


def _split_passages(text, max_chars=1200):
    """把一页文本拆成较短段落，避免把整页都塞给模型。"""
    text = clean_text(text or "")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    raw = [x.strip() for x in re.split(r"\n\s*\n", text) if x.strip()]
    if len(raw) <= 1:
        raw = [x.strip() for x in re.split(r"(?<=[.!?。！？])\s+", text) if x.strip()]
    out, buf = [], ""
    for part in raw:
        if len(part) > max_chars:
            if buf:
                out.append(buf.strip())
                buf = ""
            for i in range(0, len(part), max_chars):
                out.append(part[i:i + max_chars].strip())
        elif len(buf) + len(part) + 2 <= max_chars:
            buf = (buf + "\n" + part).strip() if buf else part
        else:
            if buf:
                out.append(buf.strip())
            buf = part
    if buf:
        out.append(buf.strip())
    return out


def retrieve_relevant_passages(book, query, top_k=6):
    """从当前 PDF 全文中检索与问题最相关的段落，并保留页码。

    这是轻量 RAG：不需要 embedding/向量数据库，优点是易部署；缺点是语义召回能力有限。
    """
    terms = _tokenize_query(query)
    if not terms:
        terms = [(query or "").lower().strip()] if query else []
    phrase = (query or "").lower().strip()
    scored = []
    try:
        pages = book_pages_text(book)
    except Exception:
        return []
    for i, page_text in enumerate(pages):
        for passage in _split_passages(page_text):
            low = passage.lower()
            score = 0
            if phrase and phrase in low:
                score += 8
            for t in terms:
                if not t:
                    continue
                c = low.count(t)
                if c:
                    score += min(c, 5)
            if re.search(r"abstract|introduction|method|methods|result|results|discussion|conclusion|finding|findings|摘要|引言|方法|结果|讨论|结论", low):
                score += 0.5
            if score > 0:
                scored.append({"page": i + 1, "score": round(score, 2), "text": passage[:1500]})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]



# -------- 轻量向量 RAG / Embedding 索引 --------
def _l2_normalize(vec):
    norm = math.sqrt(sum(float(x) * float(x) for x in vec)) or 1.0
    return [float(x) / norm for x in vec]


def _cosine(a, b):
    if not a or not b:
        return 0.0
    n = min(len(a), len(b))
    return sum(float(a[i]) * float(b[i]) for i in range(n))


def _local_hash_embedding(text, dims=384):
    """本地兜底向量：不用外部依赖。不是大模型 embedding，但能提供稳定的语义检索后备。"""
    text = (text or "").lower()
    tokens = re.findall(r"[a-zäöüß0-9_]{2,}|[\u4e00-\u9fff]", text, flags=re.I)
    grams = []
    compact = re.sub(r"\s+", "", text)
    for n in (2, 3, 4):
        grams.extend(compact[i:i+n] for i in range(max(0, len(compact)-n+1)) if compact[i:i+n].strip())
    vec = [0.0] * dims
    for tok in tokens + grams:
        h = hash(tok) % dims
        vec[h] += 1.0
    return _l2_normalize(vec)


def _embedding_client(provider, api_key, base_url=None):
    from openai import OpenAI
    if provider == "OpenAI":
        return OpenAI(api_key=api_key)
    if provider == "ZAI":
        return OpenAI(api_key=api_key, base_url=ZAI_BASE)
    if provider == "Qwen":
        return OpenAI(api_key=api_key, base_url=QWEN_BASE)
    if provider == "CustomOpenAI":
        if not base_url:
            raise ValueError("自定义 embedding 需要 Base URL")
        return OpenAI(api_key=api_key, base_url=base_url)
    raise ValueError("当前引擎不支持远程 embeddings，已使用本地轻量向量兜底")


def embed_texts(provider, api_key, texts, embed_model=None, base_url=None):
    """返回向量列表。OpenAI/Z.ai/通义千问/自定义兼容接口可调用远程 embeddings；其他情况使用本地向量。"""
    texts = [clean_text(t or "")[:8000] for t in texts]
    if provider in ("OpenAI", "ZAI", "Qwen", "CustomOpenAI") and api_key:
        if provider == "OpenAI":
            model = embed_model or "text-embedding-3-small"
        elif provider == "Qwen":
            model = embed_model or "text-embedding-v3"
        else:
            model = embed_model or "embedding-3"
        try:
            cli = _embedding_client(provider, api_key, base_url=base_url)
            out = []
            for i in range(0, len(texts), 64):
                r = cli.embeddings.create(model=model, input=texts[i:i+64])
                out.extend([_l2_normalize(item.embedding) for item in r.data])
            return out, {"provider": provider, "model": model, "kind": "remote"}
        except Exception:
            # 远程 embedding 失败时，不中断程序，改用本地向量。
            pass
    return [_local_hash_embedding(t) for t in texts], {"provider": "local", "model": "hash-384", "kind": "local"}


def rag_index_path(bid):
    safe = re.sub(r'[^\w.\-]+', '_', bid)
    return RAG_DIR / (safe + ".json")


def build_rag_index(book, provider="local", api_key="", embed_model=None, base_url=None, max_passages=1400):
    pages = book_pages_text(book)
    passages = []
    for i, txt in enumerate(pages):
        for j, passage in enumerate(_split_passages(txt, max_chars=1100)):
            if len(passage.strip()) < 40:
                continue
            passages.append({"id": f"p{i+1}_{j+1}", "page": i + 1, "text": passage[:1600]})
            if len(passages) >= max_passages:
                break
        if len(passages) >= max_passages:
            break
    if not passages:
        return {"ok": False, "error": "没有可索引文本。扫描版 PDF 需要先 OCR。"}
    vectors, meta = embed_texts(provider, api_key, [p["text"] for p in passages], embed_model=embed_model, base_url=base_url)
    for psg, vec in zip(passages, vectors):
        psg["embedding"] = vec
    idx = {"book_id": book["id"], "book_name": book.get("name", book["id"]),
           "built": datetime.now().strftime("%Y-%m-%d %H:%M"), "meta": meta, "passages": passages}
    save_json(rag_index_path(book["id"]), idx)
    return {"ok": True, "count": len(passages), "meta": meta, "built": idx["built"]}


def load_or_build_rag_index(book, provider="local", api_key="", embed_model=None, base_url=None):
    path = rag_index_path(book["id"])
    idx = load_json(path, None)
    if idx and idx.get("passages"):
        return idx
    res = build_rag_index(book, provider=provider, api_key=api_key, embed_model=embed_model, base_url=base_url)
    if not res.get("ok"):
        return None
    return load_json(path, None)


def retrieve_semantic_passages(book, query, top_k=8, provider="local", api_key="", embed_model=None, base_url=None):
    idx = load_or_build_rag_index(book, provider=provider, api_key=api_key, embed_model=embed_model, base_url=base_url)
    if not idx:
        return retrieve_relevant_passages(book, query, top_k=top_k)
    meta = idx.get("meta", {})
    q_provider = meta.get("provider", "local")
    q_model = meta.get("model")
    q_vecs, _ = embed_texts(q_provider if q_provider != "local" else "local", api_key, [query], embed_model=q_model, base_url=base_url)
    qv = q_vecs[0]
    scored = []
    for psg in idx.get("passages", []):
        score = _cosine(qv, psg.get("embedding", []))
        # 关键词小加权，避免纯向量漏掉精确术语。
        low = psg.get("text", "").lower()
        for t in _tokenize_query(query):
            if t in low:
                score += 0.03
        scored.append({"page": psg.get("page"), "score": round(float(score), 4), "text": psg.get("text", "")})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]


def detect_chapters(book):
    """基于文本行的启发式章节识别。识别不到时按每 5 页分段。"""
    pages = book_pages_text(book)
    chapters = []
    pat = re.compile(r"^(\d+(?:\.\d+)*\s+)?(abstract|introduction|background|theory|methods?|methodology|data|results?|findings?|discussion|conclusions?|references|摘要|引言|背景|理论|方法|数据|结果|讨论|结论|参考文献)\b", re.I)
    for i, txt in enumerate(pages):
        lines = [re.sub(r"\s+", " ", x).strip() for x in (txt or "").splitlines()]
        for line in lines[:35]:
            if 3 <= len(line) <= 90 and pat.search(line):
                title = line[:90]
                if not chapters or chapters[-1]["page"] != i + 1:
                    chapters.append({"title": title, "page": i + 1})
                break
    if not chapters:
        step = 5
        chapters = [{"title": f"第 {i+1}–{min(i+step, len(pages))} 页", "page": i + 1} for i in range(0, len(pages), step)]
    # 去重并计算 end_page
    out = []
    seen = set()
    for c in chapters:
        if c["page"] in seen:
            continue
        seen.add(c["page"])
        out.append(c)
    for i, c in enumerate(out):
        c["end_page"] = (out[i+1]["page"] - 1) if i + 1 < len(out) else len(pages)
        c["index"] = i
    return out


def chapter_context(book, chapter_index, max_chars=30000):
    pages = book_pages_text(book)
    chs = detect_chapters(book)
    if not chs:
        return "", None
    chapter_index = max(0, min(int(chapter_index), len(chs) - 1))
    ch = chs[chapter_index]
    blocks, used = [], 0
    for pno in range(ch["page"], ch["end_page"] + 1):
        txt = clean_text(pages[pno - 1] or "").strip()
        if not txt:
            continue
        block = f"\n\n[第{pno}页]\n{txt}"
        if used + len(block) > max_chars:
            remain = max_chars - used
            if remain > 500:
                blocks.append(block[:remain])
            break
        blocks.append(block)
        used += len(block)
    return "".join(blocks).strip(), ch


def _highlight_rgb(color):
    table = {
        "yellow": (1.0, 0.92, 0.18),
        "green": (0.35, 1.0, 0.35),
        "red": (1.0, 0.35, 0.35),
        "blue": (0.35, 0.55, 1.0),
    }
    return table.get(color, table["yellow"])


def export_annotated_pdf(book, highlights):
    doc = fitz.open(resolve_pdf_path(book))
    for h in highlights:
        try:
            page_no = int(h.get("page", 1)) - 1
            if page_no < 0 or page_no >= doc.page_count:
                continue
            pg = doc[page_no]
            rects = []
            for r in h.get("rects") or []:
                if len(r) >= 4:
                    x, y, w, hh = [float(v) for v in r[:4]]
                    rects.append(fitz.Rect(x * pg.rect.width, y * pg.rect.height,
                                           (x + w) * pg.rect.width, (y + hh) * pg.rect.height))
            # 兼容旧版纯文本重点：尽量用搜索结果生成标注。
            if not rects and h.get("text"):
                rects = pg.search_for((h.get("text") or "")[:120])[:8]
            if not rects:
                continue
            annot = pg.add_highlight_annot(rects)
            annot.set_colors(stroke=_highlight_rgb(h.get("color", "yellow")))
            annot.set_opacity(0.38)
            annot.update()
        except Exception:
            continue
    out = EXPORT_DIR / (Path(book.get("name") or book["id"]).stem + "_annotated.pdf")
    doc.save(str(out), garbage=4, deflate=True)
    doc.close()
    return out

def build_paper_card_context(book, max_chars=32000):
    """为“文献卡片/文献总结”抽取代表性页。

    优先取首页、摘要/引言/方法/结果/讨论/结论相关页和末尾页，避免长 PDF 超出模型上下文。
    """
    try:
        pages = book_pages_text(book)
    except Exception:
        return ""
    if not pages:
        return ""

    selected = []

    def add(idx):
        if 0 <= idx < len(pages) and idx not in selected:
            selected.append(idx)

    for i in range(min(3, len(pages))):
        add(i)

    patterns = [
        r"abstract|summary|zusammenfassung|摘要",
        r"introduction|einleitung|引言|导论",
        r"method|methods|methodology|data and methods|方法|数据与方法",
        r"result|results|finding|findings|结果|发现",
        r"discussion|讨论",
        r"conclusion|conclusions|结论|concluding remarks",
        r"limitation|limitations|局限",
    ]
    for pat in patterns:
        found = 0
        for i, t in enumerate(pages):
            if re.search(pat, (t or "")[:2500], flags=re.I):
                add(i)
                if i + 1 < len(pages):
                    add(i + 1)
                found += 1
                if found >= 2:
                    break

    for i in range(max(0, len(pages) - 3), len(pages)):
        add(i)

    chunks = []
    used = 0
    for idx in selected:
        txt = clean_text(pages[idx] or "").strip()
        if not txt:
            continue
        block = f"\n\n[第{idx + 1}页]\n{txt}"
        if used + len(block) > max_chars:
            remain = max_chars - used
            if remain > 500:
                chunks.append(block[:remain])
            break
        chunks.append(block)
        used += len(block)
    return "".join(chunks).strip()


from ai_clients import *  # AI 客户端层（文本/流式/视觉/SSE/引擎与 Key 解析）

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 300 * 1024 * 1024
# 仅供本机使用：禁止把 cookie/Authorization 在跨站场景下被滥用（本程序不用 cookie 登录）。
app.config["SESSION_COOKIE_SAMESITE"] = "Strict"

# 允许的本机主机名（含 IPv6 回环）。其它 Host 一律拒绝，
# 用于防御「DNS 重绑定 / 跨站请求」——即某个恶意网站把它的域名解析到 127.0.0.1，
# 再借你的浏览器偷偷访问本机这台服务器读取你的库、笔记、PDF。
_ALLOWED_HOSTS = {"127.0.0.1", "localhost", "::1", "0.0.0.0", "[::1]"}


@app.before_request
def _guard_host():
    host = (request.host or "").strip()
    # 去掉端口，只比主机名（IPv6 形如 [::1]:5000）
    hostname = host
    if hostname.startswith("["):
        hostname = hostname.split("]")[0] + "]"
    elif ":" in hostname:
        hostname = hostname.rsplit(":", 1)[0]
    if hostname and hostname not in _ALLOWED_HOSTS:
        return ("Forbidden host", 403)
    return None


@app.after_request
def add_security_headers(resp):
    # 安全加固说明：
    #  · 不下发任何允许跨域读取的 CORS 头——本程序网页与接口同源，浏览器内部调用无需 CORS；
    #    若放开 Access-Control-Allow-Origin:*，你浏览的任意网站都能跨站偷读本机接口（库/笔记/高亮/PDF 文本）。
    #  · nosniff 防止浏览器把响应猜成可执行类型；DENY 禁止本页被任何站点用 iframe 套住（点击劫持）。
    #  · no-referrer 避免把本机地址带到外部。
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-Frame-Options"] = "DENY"
    resp.headers["Referrer-Policy"] = "no-referrer"
    return resp


@app.errorhandler(Exception)
def handle_unexpected_error(e):
    try:
        from werkzeug.exceptions import HTTPException
        if isinstance(e, HTTPException):
            return e
    except Exception:
        pass
    append_log("Flask 未处理异常", traceback.format_exc())
    return jsonify({"error": "服务器内部错误，详情已写入 data/server_error.log"}), 500




REQUIRED_MODULES = {
    "Flask Web Server": ("flask", "flask"),
    "PyMuPDF PDF Reader": ("pymupdf", "fitz"),
    "OpenAI / OpenAI-compatible SDK": ("openai", "openai"),
    "Anthropic Claude SDK": ("anthropic", "anthropic"),
    "Free Google Translate Helper": ("deep-translator", "deep_translator"),
    "Word Export": ("python-docx", "docx"),
    "PDF Export": ("reportlab", "reportlab"),
    "Web Article Fetch": ("trafilatura", "trafilatura"),
    "Production Local Server": ("waitress", "waitress"),
}


def _check_module(mod_name):
    try:
        __import__(mod_name)
        return True, "ok"
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def _check_writable(path: Path):
    try:
        path.mkdir(parents=True, exist_ok=True)
        test = path / "._apw_write_test.tmp"
        test.write_text("ok", encoding="utf-8")
        test.unlink(missing_ok=True)
        return True, "ok"
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"



@app.route("/")
def index():
    pdfjs_local = (LIB_DIR / "pdf.min.js").exists() or (BUNDLED_LIB_DIR / "pdf.min.js").exists()
    worker_local = (LIB_DIR / "pdf.worker.min.js").exists() or (BUNDLED_LIB_DIR / "pdf.worker.min.js").exists()
    pdfjs = "/lib/pdf.min.js" if pdfjs_local else PDFJS_CDN
    worker = "/lib/pdf.worker.min.js" if worker_local else WORKER_CDN
    html = (WEB_DIR / "index.html").read_text(encoding="utf-8")
    html = html.replace("__PDFJS__", pdfjs).replace("__WORKER__", worker)
    # 缓存破坏：给本地 CSS/JS 附上基于文件修改时间的版本号，确保前端改动后
    # 浏览器（含内嵌 WebView2）一定加载最新文件，而不是旧的启发式缓存副本。
    try:
        css_v = int((WEB_DIR / "style.css").stat().st_mtime)
        js_v = int((WEB_DIR / "app.js").stat().st_mtime)
    except OSError:
        css_v = js_v = 0
    html = (html.replace('/web/style.css"', '/web/style.css?v=%d"' % css_v)
                .replace('/web/app.js"', '/web/app.js?v=%d"' % js_v))
    resp = Response(html, mimetype="text/html")
    resp.headers["Cache-Control"] = "no-store"
    return resp


@app.route("/web/<path:fname>")
def web_static(fname):
    # 前端静态资源（index.html / style.css / app.js）。同样做路径穿越校验。
    base_r = WEB_DIR.resolve()
    try:
        pp = (base_r / fname).resolve()
    except Exception:
        return "not found", 404
    if not pp.is_relative_to(base_r) or not (pp.exists() and pp.is_file()):
        return "not found", 404
    resp = send_file(str(pp))
    resp.headers["Cache-Control"] = "no-store"  # 禁止浏览器缓存前端静态文件，避免改动后仍加载旧副本
    return resp


@app.route("/lib/<path:fname>")
def api_lib(fname):
    # 安全加固：校验解析后的真实路径必须落在 lib 目录内，阻止 ../ 路径穿越
    # （否则可被 /lib/../data/notes.json 之类请求读到笔记、库、PDF 乃至任意本地文件）。
    for base in (LIB_DIR, BUNDLED_LIB_DIR):
        try:
            base_r = base.resolve()
            p = (base_r / fname).resolve()
        except Exception:
            continue
        if not p.is_relative_to(base_r):
            continue
        if p.exists() and p.is_file():
            return send_file(str(p))
    return "not found", 404

@app.route("/api/library")
def api_library():
    return jsonify(safe_library())


@app.route("/api/lib/set_group", methods=["POST"])
def api_lib_set_group():
    """把某本书归入某个分组（二级书架）。group 留空表示移出分组（归到“未分类”）。"""
    d = request.get_json(force=True) or {}
    bid = str(d.get("id") or "").strip()
    group = str(d.get("group") or "").strip()[:60]
    if not bid:
        return jsonify({"error": "缺少书目 id"}), 400
    lib = safe_library()
    hit = False
    for b in lib:
        if b.get("id") == bid:
            if group:
                b["group"] = group
            else:
                b.pop("group", None)
            hit = True
            break
    if not hit:
        return jsonify({"error": "未找到该书目"}), 404
    save_json(LIB_FILE, lib)
    return jsonify({"ok": True, "id": bid, "group": group})


def _within_pdf_dir(path):
    """安全：只允许删除位于 data/pdfs/ 目录内的文件，杜绝路径穿越误删其它文件。"""
    try:
        p = Path(path).resolve()
        root = PDF_DIR.resolve()
        return str(p).startswith(str(root) + os.sep) and p.is_file()
    except Exception:
        return False


@app.route("/api/lib/delete", methods=["POST"])
def api_lib_delete():
    """删除一本书：移出书库、删除其 PDF 物理文件（仅限 data/pdfs 内）、
    并清理它的高亮 / 笔记 / RAG 索引 / 文本缓存，避免留下孤儿数据。"""
    d = request.get_json(force=True) or {}
    bid = str(d.get("id") or "").strip()
    if not bid:
        return jsonify({"error": "缺少书目 id"}), 400
    lib = safe_library()
    target = None
    rest = []
    for b in lib:
        if b.get("id") == bid and target is None:
            target = b
        else:
            rest.append(b)
    if target is None:
        return jsonify({"error": "未找到该书目"}), 404

    # 1) 删除物理 PDF（仅当它确实在 data/pdfs 内）
    removed_file = False
    try:
        real = resolve_pdf_path(target)
        if real and _within_pdf_dir(real):
            Path(real).unlink(missing_ok=True)
            removed_file = True
        else:
            cand = PDF_DIR / bid
            if cand.exists() and _within_pdf_dir(str(cand)):
                cand.unlink(missing_ok=True)
                removed_file = True
    except Exception:
        append_log("删除 PDF 物理文件失败", traceback.format_exc())

    # 2) 写回书库
    save_json(LIB_FILE, rest)

    # 3) 清理高亮 / 笔记 / RAG 索引 / 文本缓存
    try:
        hls = load_json(HL_FILE, [])
        hls2 = [h for h in hls if h.get("book_id") != bid]
        if len(hls2) != len(hls):
            save_json(HL_FILE, hls2)
    except Exception:
        pass
    try:
        notes = load_json(NOTE_FILE, [])
        if isinstance(notes, dict):
            if bid in notes:
                notes.pop(bid, None)
                save_json(NOTE_FILE, notes)
        elif isinstance(notes, list):
            n2 = [n for n in notes if n.get("book_id") != bid]
            if len(n2) != len(notes):
                save_json(NOTE_FILE, n2)
    except Exception:
        pass
    try:
        rp = rag_index_path(bid)
        if rp.exists():
            rp.unlink(missing_ok=True)
    except Exception:
        pass
    _text_cache.pop(bid, None)
    return jsonify({"ok": True, "id": bid, "removed_file": removed_file})


@app.route("/api/lib/rename_group", methods=["POST"])
def api_lib_rename_group():
    """重命名一个分组：把所有属于旧分组名的书改到新分组名。"""
    d = request.get_json(force=True) or {}
    old = str(d.get("old") or "").strip()
    new = str(d.get("new") or "").strip()[:60]
    if not old or not new:
        return jsonify({"error": "缺少分组名"}), 400
    lib = safe_library()
    n = 0
    for b in lib:
        if (b.get("group") or "").strip() == old:
            b["group"] = new
            n += 1
    save_json(LIB_FILE, lib)
    return jsonify({"ok": True, "moved": n, "new": new})


@app.route("/api/lib/delete_group", methods=["POST"])
def api_lib_delete_group():
    """删除一个分组（仅解散分组，不删书）：把该分组下的书移出到“未分类”。"""
    d = request.get_json(force=True) or {}
    name = str(d.get("name") or "").strip()
    if not name:
        return jsonify({"error": "缺少分组名"}), 400
    lib = safe_library()
    n = 0
    for b in lib:
        if (b.get("group") or "").strip() == name:
            b.pop("group", None)
            n += 1
    save_json(LIB_FILE, lib)
    return jsonify({"ok": True, "released": n})


# ---------------- AI 内容集合（保存各 AI 功能产出，像笔记本一样可查看/编辑） ----------------
@app.route("/api/ai_collection")
def api_ai_collection():
    items = safe_ai_collection()
    # 新的在前
    items.sort(key=lambda x: x.get("time", ""), reverse=True)
    return jsonify(items)


@app.route("/api/ai_collection/save", methods=["POST"])
def api_ai_collection_save():
    try:
        d = request.get_json(force=True, silent=True) or {}
        content = (d.get("content") or "").strip()
        if not content:
            return jsonify({"error": "没有可保存的内容"}), 400
        category = (d.get("category") or "其他").strip() or "其他"
        title = (d.get("title") or "").strip()
        if not title:
            title = content.splitlines()[0][:40] if content else "未命名"
        items = safe_ai_collection()
        item = {
            "id": "ai" + uuid.uuid4().hex[:12],
            "category": category,
            "title": title,
            "content": content,
            "book": (d.get("book") or "").strip(),
            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        items.append(item)
        save_json(AI_COLLECTION_FILE, items)
        return jsonify(item)
    except Exception:
        append_log("/api/ai_collection/save 异常", traceback.format_exc())
        return jsonify({"error": "保存失败"}), 500


@app.route("/api/ai_collection/update", methods=["POST"])
def api_ai_collection_update():
    """允许用户在集合窗口里修改已保存内容（标题 / 正文）。"""
    try:
        d = request.get_json(force=True, silent=True) or {}
        cid = d.get("id")
        items = safe_ai_collection()
        hit = None
        for it in items:
            if it.get("id") == cid:
                hit = it
                break
        if not hit:
            return jsonify({"error": "未找到该条目"}), 404
        if "title" in d:
            hit["title"] = (d.get("title") or "").strip() or hit.get("title") or "未命名"
        if "content" in d:
            hit["content"] = d.get("content") or ""
        save_json(AI_COLLECTION_FILE, items)
        return jsonify(hit)
    except Exception:
        append_log("/api/ai_collection/update 异常", traceback.format_exc())
        return jsonify({"error": "更新失败"}), 500


@app.route("/api/ai_collection/delete", methods=["POST"])
def api_ai_collection_delete():
    try:
        d = request.get_json(force=True, silent=True) or {}
        cid = d.get("id")
        items = [it for it in safe_ai_collection() if it.get("id") != cid]
        save_json(AI_COLLECTION_FILE, items)
        return jsonify({"ok": True})
    except Exception:
        append_log("/api/ai_collection/delete 异常", traceback.format_exc())
        return jsonify({"error": "删除失败"}), 500


@app.route("/api/upload", methods=["POST"])
def api_upload():
    """导入 PDF。
    修复点：
    1) 不再覆盖同名文件，避免 Windows 中文件被浏览器/fitz 占用导致 HTTP 500；
    2) 自动修复旧版/损坏 library.json；
    3) 所有异常写入 data/server_error.log，并向前端返回可读错误。
    """
    try:
        f = request.files.get("file")
        if not f or not getattr(f, "filename", ""):
            return jsonify({"error": "没有选择 PDF 文件"}), 400
        original_name = f.filename
        if not original_name.lower().endswith(".pdf"):
            return jsonify({"error": "目前只支持导入 PDF 文件"}), 400

        safe_id, dest = unique_pdf_dest(original_name)
        try:
            f.save(str(dest))
        except Exception as e:
            append_log("PDF 保存失败", traceback.format_exc())
            return jsonify({"error": f"保存 PDF 失败：{e}"}), 500

        try:
            doc = fitz.open(str(dest))
            pc = doc.page_count
            if pc <= 0:
                raise ValueError("PDF 页数为 0")
            doc.close()
        except Exception as e:
            try:
                dest.unlink(missing_ok=True)
            except Exception:
                pass
            append_log("PDF 读取失败", traceback.format_exc())
            return jsonify({"error": f"无法读取 PDF，可能文件损坏或不是标准 PDF：{e}"}), 400

        lib = safe_library()
        book = {
            "id": safe_id,
            "name": original_name,
            "path": str(dest),
            "total_pages": pc,
            "current_page": 0,
            "added": datetime.now().strftime("%Y-%m-%d"),
        }
        lib.append(book)
        save_json(LIB_FILE, lib)
        _text_cache.pop(safe_id, None)
        return jsonify(book)
    except Exception:
        append_log("/api/upload 未处理异常", traceback.format_exc())
        return jsonify({"error": "服务器导入时发生未处理错误；详情已写入 data/server_error.log"}), 500


@app.route("/api/pdf/<bid>")
def api_pdf(bid):
    b = find_book(bid)
    if not b:
        return "not found", 404
    path = resolve_pdf_path(b)
    if not path or not Path(path).exists():
        append_log("PDF 文件缺失", f"id={bid} 记录路径={b.get('path')!r} 实际未找到")
        return jsonify({"error": "PDF 文件已不在原位置（可能被移动或删除）。请在书架重新导入该 PDF。"}), 404
    return send_file(path, mimetype="application/pdf")


@app.route("/api/pagetext/<bid>")
def api_pagetext(bid):
    b = find_book(bid)
    if not b:
        return jsonify({"text": ""})
    try:
        idx = int(request.args.get("page", "1")) - 1
        pages = book_pages_text(b)
        text = pages[idx] if 0 <= idx < len(pages) else ""
        if not (text or "").strip():
            cached = safe_ocr_cache().get(f"{bid}:{idx+1}")
            if isinstance(cached, dict):
                text = cached.get("text") or ""
        return jsonify({"text": text})
    except Exception:
        return jsonify({"text": ""})




@app.route("/api/progress", methods=["POST"])
def api_progress():
    d = request.get_json(force=True)
    lib = load_json(LIB_FILE, [])
    for b in lib:
        if b["id"] == d["id"]:
            b["current_page"] = int(d["page"])
    save_json(LIB_FILE, lib)
    return jsonify({"ok": True})


@app.route("/api/highlights/<bid>")
def api_get_hl(bid):
    return jsonify([h for h in safe_highlights() if h.get("book_id") == bid])


@app.route("/api/highlights", methods=["POST"])
def api_add_hl():
    d = request.get_json(force=True)
    hls = safe_highlights()
    item = {
        "id": d.get("id") or ("h" + uuid.uuid4().hex[:12]),
        "book_id": d["book_id"],
        "page": int(d["page"]),
        "text": d.get("text", ""),
        "color": d.get("color", "yellow"),
        "rects": d.get("rects") or [],
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    hls.append(item)
    save_json(HL_FILE, hls)
    return jsonify({"ok": True, "highlight": item})


@app.route("/api/highlights/delete", methods=["POST"])
def api_del_hl():
    d = request.get_json(force=True)
    hid = d.get("id")
    if hid:
        hls = [h for h in safe_highlights() if h.get("id") != hid]
    else:
        hls = [h for h in safe_highlights()
               if not (h.get("book_id") == d.get("book_id") and h.get("text") == d.get("text")
                       and h.get("time") == d.get("time"))]
    save_json(HL_FILE, hls)
    return jsonify({"ok": True})


@app.route("/api/notes/<bid>")
def api_get_notes(bid):
    return jsonify(notes_list(load_json(NOTE_FILE, {}), bid))


@app.route("/api/notes/save", methods=["POST"])
def api_save_note():
    d = request.get_json(force=True)
    alln = load_json(NOTE_FILE, {})
    lst = notes_list(alln, d["book_id"])
    note = d["note"]
    found = False
    for n in lst:
        if n["id"] == note["id"]:
            n["title"] = note["title"]
            n["body"] = note["body"]
            found = True
    if not found:
        lst.append(note)
    alln[d["book_id"]] = lst
    save_json(NOTE_FILE, alln)
    return jsonify({"ok": True})


@app.route("/api/notes/delete", methods=["POST"])
def api_del_note():
    d = request.get_json(force=True)
    alln = load_json(NOTE_FILE, {})
    alln[d["book_id"]] = [n for n in notes_list(alln, d["book_id"]) if n["id"] != d["id"]]
    save_json(NOTE_FILE, alln)
    return jsonify({"ok": True})


@app.route("/api/search/<bid>")
def api_search(bid):
    q = (request.args.get("q") or "").strip().lower()
    b = find_book(bid)
    if not b or not q:
        return jsonify({"results": [], "stop": q in STOPWORDS})
    res = []
    for i, t in enumerate(book_pages_text(b)):
        idx = t.lower().find(q)
        if idx != -1:
            s, e = max(0, idx - 60), min(len(t), idx + len(q) + 60)
            res.append({"page": i + 1, "snippet": t[s:e].replace("\n", " ").strip()})
    return jsonify({"results": res[:80], "stop": q in STOPWORDS})


@app.route("/api/translate", methods=["POST"])
def api_translate():
    d = request.get_json(force=True)
    text = clean_text((d.get("text") or "").strip())
    if not text:
        return jsonify({"error": "没有文本"}), 400
    target = d.get("target", "简体中文")
    if d.get("provider") == "免费":
        try:
            from deep_translator import GoogleTranslator
            code = {"简体中文": "zh-CN", "English": "en", "Deutsch": "de"}.get(target, "zh-CN")
            out = [GoogleTranslator(source="auto", target=code).translate(text[i:i + 4500])
                   for i in range(0, len(text), 4500)]
            return jsonify({"result": "\n".join(x for x in out if x)})
        except Exception as e:
            return jsonify({"error": f"免费翻译失败：{e}"}), 500
    provider, key = resolve_key(d)
    if not key:
        return jsonify({"error": "请先在设置里填入 API Key"}), 400
    sys = (f"你是专业的学术译者。请把用户文本忠实、流畅地翻译为{target}，"
           "保留段落结构，只输出译文，不要任何解释。")
    sys = with_glossary(sys, d.get("glossary"))
    try:
        return jsonify({"result": llm_call(provider, key, sys, text, model=d.get("model"), effort=d.get("effort"), base_url=d.get("base_url"))})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/page_labels/<bid>")
def api_page_labels(bid):
    b = find_book(bid)
    if not b:
        return jsonify({"error": "not found"}), 404
    return jsonify({"labels": _page_labels(b), "offset": int(b.get("page_offset") or 0)})


@app.route("/api/page_offset", methods=["POST"])
def api_page_offset():
    d = request.get_json(force=True)
    b = find_book(d.get("id"))
    if not b:
        return jsonify({"error": "not found"}), 404
    off = int(d.get("offset") or 0)
    lib = safe_library()
    for x in lib:
        if x.get("id") == b.get("id"):
            x["page_offset"] = off
    save_json(LIB_FILE, lib)
    _PLABEL_CACHE.pop(b.get("id"), None)
    return jsonify({"ok": True, "offset": off, "labels": _page_labels(find_book(b.get("id")))})


@app.route("/api/explain_stream", methods=["POST"])
def api_explain_stream():
    """AI 解读·解释（流式）：三块——句子翻译 / 关键术语 / 结合上下文的含义。
    上下文范围 ctx_scope：around=前后页(默认) / all=全文检索 / range=自定义页码。
    浅白中文、术语保持专业、以原文语言自身语义为准；思考与否由 effort 控制（前端的“思考”开关）。"""
    d = request.get_json(force=True)
    text = clean_text((d.get("text") or "").strip())
    if not text:
        return jsonify({"error": "没有要解释的文字"}), 400
    provider, key = resolve_key(d)
    if not key:
        return jsonify({"error": "请先在设置里填入当前引擎的 API Key"}), 400
    scope = d.get("ctx_scope") or "around"
    ctx_note = ""
    b = find_book(d.get("book_id")) if d.get("book_id") else None
    if b:
        try:
            pages = book_pages_text(b)
            n = len(pages)
            if scope == "all":
                hits = retrieve_semantic_passages(b, text, top_k=6, provider=provider, api_key=key,
                                                  embed_model=d.get("embed_model"), base_url=d.get("base_url"))
                ctx_note = "\n\n".join(f"[第{_plab(b, h['page'])}页]\n{h['text']}" for h in hits)
            elif scope == "range":
                sel = _parse_pages(d.get("page_range") or "", n)
                ctx_note = "\n\n".join(f"[第{_plab(b, p)}页]\n{clean_text(pages[p - 1] or '').strip()}" for p in sel)
            else:  # around：前后页（当前页 ±1）
                pno = max(1, min(int(d.get("page") or 1), n))
                lo, hi = max(1, pno - 1), min(n, pno + 1)
                ctx_note = "\n\n".join(f"[第{_plab(b, p)}页]\n{clean_text(pages[p - 1] or '').strip()}"
                                       for p in range(lo, hi + 1))
            ctx_note = ctx_note[:6000]
        except Exception:
            ctx_note = ""
    sys = ("你是耐心精准的学术阅读助手，帮读者读懂一段外文（多为德语/英语）学术文本。"
           "用【浅白易懂的中文和简单句式】解释，但【专业术语保持准确】，必要时保留原文词并附通行译名。"
           "【以原文语言自身的语义与用法为准】去解释德/英概念，不要用中文语境的现成框架去套；"
           "忠于原文、不臆测，不确定处标注(不确定)；最终全部用简体中文。" + IRON)
    sys = with_glossary(sys, d.get("glossary"))
    user = ("请针对下面【选中的文字】给出三块，简短为要：\n"
            "① 句子翻译：先给一句忠实、通顺的中文翻译；\n"
            "② 关键术语：挑出最关键的 3–5 个，每个用一句浅白的话讲清它在原文语言里的含义与用法，保留原文词；\n"
            "③ 上下文含义：结合下面的【上下文】，说清这句/这段在此处到底在讲什么、承接或转折了什么。\n"
            "浅白但术语专业，不要长篇大论。\n\n【选中的文字】\n" + text
            + (("\n\n【上下文（前后页 / 检索）】\n" + ctx_note) if ctx_note else ""))
    return sse_stream(llm_chat_stream(provider, key, sys, [{"role": "user", "content": user}],
                                      model=d.get("model"), effort=d.get("effort"),
                                      base_url=d.get("base_url")))


@app.route("/api/ask", methods=["POST"])
def api_ask():
    d = request.get_json(force=True)
    q = (d.get("question") or "").strip()
    if not q:
        return jsonify({"error": "没有问题"}), 400
    provider, key = resolve_key(d)
    if not key:
        return jsonify({"error": "请先在设置里填入当前引擎的 API Key"}), 400
    ctx = clean_text(d.get("context", ""))
    rag_note = ""
    if d.get("use_rag") and d.get("book_id"):
        b = find_book(d.get("book_id"))
        if b:
            hits = retrieve_relevant_passages(b, q, top_k=6)
            rag_note = "\n\n【系统从全文检索到的相关段落】\n" + "\n\n".join(
                f"[第{_plab(b, h['page'])}页 | score={h['score']}]\n{h['text']}" for h in hits
            )
    history = d.get("history", []) or []
    sys = ("你是学术阅读导师。请用简体中文回答用户的问题。"
           "回答时优先依据材料，不要凭空编造；如果依据来自全文检索段落，请在关键判断后标注页码，例如（第12页）。"
           "如果材料不足以回答，请明确说明材料不足。\n\n【用户选中的材料 / 笔记】\n" + ctx + rag_note)
    sys = with_glossary(sys, d.get("glossary"))
    msgs = [{"role": m["role"], "content": m["content"]}
            for m in history if m.get("role") in ("user", "assistant")]
    msgs.append({"role": "user", "content": q})
    try:
        return jsonify({"result": llm_chat(provider, key, sys, msgs,
                                           model=d.get("model"), effort=d.get("effort"),
                                           base_url=d.get("base_url"))})
    except Exception as e:
        return jsonify({"error": str(e)}), 500





@app.route("/api/build_rag_index", methods=["POST"])
def api_build_rag_index():
    d = request.get_json(force=True)
    bid = d.get("book_id")
    b = find_book(bid) if bid else None
    if not b:
        return jsonify({"error": "请先打开一篇 PDF"}), 400
    provider, key = resolve_key(d)
    # 没有 key 或当前引擎不支持 embedding 时，自动使用本地轻量向量。
    use_provider = provider if key else "local"
    try:
        res = build_rag_index(b, provider=use_provider, api_key=key,
                              embed_model=d.get("embed_model"), base_url=d.get("base_url"))
        if not res.get("ok"):
            return jsonify({"error": res.get("error", "索引失败")}), 400
        return jsonify(res)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/rag_status/<bid>")
def api_rag_status(bid):
    path = rag_index_path(bid)
    idx = load_json(path, None)
    if not idx:
        return jsonify({"exists": False})
    return jsonify({"exists": True, "count": len(idx.get("passages", [])),
                    "built": idx.get("built"), "meta": idx.get("meta", {})})


@app.route("/api/ask_stream", methods=["POST"])
def api_ask_stream():
    d = request.get_json(force=True)
    q = (d.get("question") or "").strip()
    if not q:
        return jsonify({"error": "没有问题"}), 400
    provider, key = resolve_key(d)
    if not key:
        return jsonify({"error": "请先在设置里填入当前引擎的 API Key"}), 400
    ctx = clean_text(d.get("context", ""))
    rag_note = ""
    if d.get("use_rag") and d.get("book_id"):
        b = find_book(d.get("book_id"))
        if b:
            hits = retrieve_semantic_passages(b, q, top_k=8, provider=provider, api_key=key,
                                             embed_model=d.get("embed_model"), base_url=d.get("base_url"))
            rag_note = "\n\n【系统从全文语义索引检索到的相关段落】\n" + "\n\n".join(
                f"[第{_plab(b, h['page'])}页 | semantic_score={h['score']}]\n{h['text']}" for h in hits
            )
    history = d.get("history", []) or []
    sys = ("你是学术阅读导师。请用简体中文回答用户的问题。"
           "回答时优先依据材料，不要凭空编造；如果依据来自全文检索段落，请在关键判断后标注页码，例如（第12页）。"
           "如果材料不足以回答，请明确说明材料不足。\n\n【用户选中的材料 / 笔记】\n" + ctx + rag_note)
    sys = with_glossary(sys, d.get("glossary"))
    msgs = [{"role": m["role"], "content": m["content"]}
            for m in history if m.get("role") in ("user", "assistant")]
    msgs.append({"role": "user", "content": q})
    return sse_stream(llm_chat_stream(provider, key, sys, msgs,
                                      model=d.get("model"), effort=d.get("effort"),
                                      base_url=d.get("base_url")))












@app.route("/api/highlights/export_pdf/<bid>")
def api_export_annotated_pdf(bid):
    b = find_book(bid)
    if not b:
        return "not found", 404
    hls = [h for h in load_json(HL_FILE, []) if h.get("book_id") == bid]
    try:
        out = export_annotated_pdf(b, hls)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return send_file(str(out), as_attachment=True, download_name=(Path(b.get("name") or bid).stem + "_annotated.pdf"))


@app.route("/api/ocr", methods=["POST"])
def api_ocr():
    d = request.get_json(force=True)
    b64 = d.get("image", "") or ""
    if b64.startswith("data:") and "," in b64:
        b64 = b64.split(",", 1)[1]
    if not b64:
        return jsonify({"error": "没有图片"}), 400
    # 安全：限制单张截图大小（base64 约为原图 1.33 倍），上限约 12MB 原图，防滥用/超大请求。
    if len(b64) > 16 * 1024 * 1024:
        return jsonify({"error": "截图区域过大，请框选更小的区域后重试"}), 400
    task = d.get("task") or "text"
    if task not in OCR_TASK_PROMPTS:
        task = "text"
    provider, key = resolve_key(d)
    if provider == "免费" or not key:
        return jsonify({"error": "框选 OCR 需要支持图片输入的模型并填好 key（建议用 通义千问 / 智谱GLM / Kimi / OpenAI / Claude）"}), 400
    try:
        return jsonify({"result": vision_ocr(provider, key, d.get("model"), b64, effort=d.get("effort"), base_url=d.get("base_url"), task=task)})
    except Exception as e:
        return jsonify({"error": f"识别失败：{e}（若当前模型不支持图片，请切到 通义千问/智谱GLM/Kimi/OpenAI/Claude）"}), 500




def _is_safe_public_url(url):
    """安全加固：只允许抓取公网地址，拒绝指向内网/本机/链路本地的 URL，
    防止本机服务器被用作 SSRF 跳板去访问内网资源或本机其它端口。"""
    import ipaddress
    import socket
    from urllib.parse import urlparse
    try:
        host = urlparse(url).hostname
        if not host:
            return False
        for info in socket.getaddrinfo(host, None):
            ip = ipaddress.ip_address(info[4][0])
            if (ip.is_private or ip.is_loopback or ip.is_link_local
                    or ip.is_reserved or ip.is_multicast or ip.is_unspecified):
                return False
        return True
    except Exception:
        return False


class _SafeRedirect(urllib.request.HTTPRedirectHandler):
    """跟随跳转时，对每一个跳转目标都重新做 SSRF 校验，
    防止『公网地址 302 跳到内网/本机地址』绕过安全检查。"""
    def redirect_request(self, req, fp, code, msg, headers, newurl):
        if not _is_safe_public_url(newurl):
            raise urllib.error.URLError("出于安全考虑，拒绝跳转到非公网地址")
        return super().redirect_request(req, fp, code, msg, headers, newurl)


def _http_get_html(url, max_bytes=5 * 1024 * 1024, timeout=20):
    """用标准库抓取网页 HTML：带浏览器 UA（很多站点没有 UA 会直接返回空），
    处理 gzip/deflate 与字符编码，限制最大体积，且跳转受 SSRF 校验保护。
    返回 (html_text, final_url)。无需任何第三方库。"""
    import gzip
    import zlib
    headers = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/124.0 Safari/537.36"),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8,de;q=0.7",
        "Accept-Encoding": "gzip, deflate",
    }
    opener = urllib.request.build_opener(_SafeRedirect())
    req = urllib.request.Request(url, headers=headers)
    with opener.open(req, timeout=timeout) as resp:
        final_url = resp.geturl()
        ctype = (resp.headers.get("Content-Type") or "").lower()
        if ctype and ("html" not in ctype and "xml" not in ctype and "text" not in ctype):
            raise ValueError("这个链接不是网页（可能是文件或图片），无法转成文章")
        raw = resp.read(max_bytes + 1)
        if len(raw) > max_bytes:
            raw = raw[:max_bytes]
        enc = resp.headers.get_content_charset()
        data = raw
        ce = (resp.headers.get("Content-Encoding") or "").lower()
        try:
            if "gzip" in ce:
                data = gzip.decompress(raw)
            elif "deflate" in ce:
                try:
                    data = zlib.decompress(raw)
                except Exception:
                    data = zlib.decompress(raw, -zlib.MAX_WBITS)
        except Exception:
            data = raw
    if not enc:
        head = data[:2048].decode("ascii", "ignore").lower()
        m = re.search(r'charset=["\']?([a-z0-9_\-]+)', head)
        enc = m.group(1) if m else "utf-8"
    try:
        html = data.decode(enc, "replace")
    except Exception:
        html = data.decode("utf-8", "replace")
    return html, final_url


class _ReadableHTML(_html_parser.HTMLParser):
    """极简正文抽取：丢弃 script/style/nav/header/footer/aside/form 等非正文区域，
    在块级元素边界插入换行，累积可读文本。不依赖 lxml / bs4。"""
    _SKIP = {"script", "style", "noscript", "head", "nav", "header", "footer",
             "aside", "form", "button", "svg", "iframe", "template"}
    _BLOCK = {"p", "div", "section", "article", "br", "li", "ul", "ol", "tr",
              "table", "blockquote", "pre", "figure", "figcaption",
              "h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []
        self._skip_depth = 0
        self.title = ""
        self._in_title = False

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1
        if tag == "title":
            self._in_title = True
        if tag in self._BLOCK and self._skip_depth == 0:
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip_depth > 0:
            self._skip_depth -= 1
        if tag == "title":
            self._in_title = False
        if tag in self._BLOCK and self._skip_depth == 0:
            self.parts.append("\n")

    def handle_data(self, data):
        if self._in_title and not self.title:
            self.title = data.strip()
        if self._skip_depth == 0 and data:
            self.parts.append(data)

    def get_text(self):
        raw = "".join(self.parts)
        lines = []
        prev = None
        for ln in raw.split("\n"):
            ln = re.sub(r"[ \t\u00a0]+", " ", ln).strip()
            if not ln:
                if lines and lines[-1] != "":
                    lines.append("")
                continue
            if ln == prev:        # 去掉相邻重复（常见于导航重复链接）
                continue
            lines.append(ln)
            prev = ln
        # 收尾：压缩多余空行
        out, blank = [], 0
        for ln in lines:
            if ln == "":
                blank += 1
                if blank > 1:
                    continue
            else:
                blank = 0
            out.append(ln)
        return "\n".join(out).strip()


def _html_to_article(html):
    """从 HTML 抽取 (title, text)。优先用 trafilatura（若已安装，质量更好），
    否则用内置的标准库抽取器——保证『没装任何额外库也能抓取成功』。"""
    title, text = "", ""
    try:
        import trafilatura  # 可选增强：装了更好，没装也不影响
        t = trafilatura.extract(html, include_comments=False,
                                 include_tables=True, favor_recall=True)
        if t and t.strip():
            text = t.strip()
        try:
            md = trafilatura.extract_metadata(html)
            if md and getattr(md, "title", None):
                title = md.title
        except Exception:
            pass
    except Exception:
        pass
    if not text:
        p = _ReadableHTML()
        try:
            p.feed(html)
        except Exception:
            pass
        text = p.get_text()
        title = title or p.title
    return (title or "网页文章"), text


@app.route("/api/fetch_url", methods=["POST"])
def api_fetch_url():
    d = request.get_json(force=True)
    url = (d.get("url") or "").strip()
    if not url.startswith("http"):
        return jsonify({"error": "请输入有效的网址（以 http 开头）"}), 400
    if not _is_safe_public_url(url):
        return jsonify({"error": "出于安全考虑，只能抓取公网网址，不允许内网/本机地址。"}), 400
    try:
        html, final_url = _http_get_html(url)
    except urllib.error.HTTPError as e:
        return jsonify({"error": f"打不开这个网页（HTTP {e.code}；可能需要登录或被反爬挡住）"}), 400
    except Exception as e:
        return jsonify({"error": f"抓取失败：{e}"}), 400
    if not _is_safe_public_url(final_url):   # 双保险：最终落地地址也必须是公网
        return jsonify({"error": "出于安全考虑，拒绝抓取该地址。"}), 400
    title, text = _html_to_article(html)
    if not text or not text.strip():
        return jsonify({"error": "没能从这个网页提取到正文（可能是纯图片页、或正文由脚本动态加载）。"}), 400
    safe = (re.sub(r'[\\/:*?"<>|\n\r\t]', "_", title).strip()[:80]) or "网页文章"
    dest = PDF_DIR / (safe + ".pdf")
    i = 1
    while dest.exists():
        dest = PDF_DIR / (safe + "_" + str(i) + ".pdf")
        i += 1
    try:
        build_article_pdf(title, text, url, str(dest))
    except Exception as e:
        append_log("网页转 PDF 失败", traceback.format_exc())
        return jsonify({"error": f"生成 PDF 失败：{e}"}), 500
    try:
        doc = fitz.open(str(dest))
        pc = doc.page_count
        doc.close()
    except Exception:
        pc = 1
    lib = load_json(LIB_FILE, [])
    bid = dest.name
    book = {"id": bid, "name": title, "path": str(dest), "total_pages": pc,
            "current_page": 0, "added": datetime.now().strftime("%Y-%m-%d")}
    lib.append(book)
    save_json(LIB_FILE, lib)
    _text_cache.pop(bid, None)
    return jsonify(book)


# -------- 生成 PDF / Word --------
def _esc(t):
    return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_article_pdf(title, text, source_url, out_path):
    """把网页正文排版成 PDF。优先用 PyMuPDF（核心依赖，必有；内置中日韩字体，
    无需任何外部字体文件），从而保证『不装 reportlab 也能成功生成』。"""
    try:
        _build_article_pdf_fitz(title, text, source_url, out_path)
        return
    except Exception:
        append_log("PyMuPDF 生成网页 PDF 失败，尝试 reportlab", traceback.format_exc())
    # 兜底：装了 reportlab 时用它（输出更精细），没装则上面的 PyMuPDF 已处理。
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    blob = (title or "") + (text or "")
    has_cjk = bool(re.search(r'[\u4e00-\u9fff]', blob))
    if has_cjk:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font = "STSong-Light"
    else:
        font = "Helvetica"
    base = ParagraphStyle("a", fontName=font, fontSize=11.5, leading=18)
    meta = ParagraphStyle("m", fontName=font, fontSize=9, leading=13, textColor="#888888")
    h1 = ParagraphStyle("h", fontName=font, fontSize=18, leading=24, spaceAfter=12)
    doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    flow = [Paragraph(_esc(title or "网页文章"), h1),
            Paragraph(_esc(source_url), meta), Spacer(1, 12)]
    for para in (text or "").split("\n"):
        para = para.strip()
        if para:
            flow.append(Paragraph(_esc(para), base))
            flow.append(Spacer(1, 6))
    doc.build(flow)


def _build_article_pdf_fitz(title, text, source_url, out_path):
    """用 PyMuPDF 把标题 + 出处 + 正文段落分页排版成 A4 PDF。
    中文用内置 china-s 字体，纯西文用 Helvetica；中西混排统一用 china-s（含拉丁字形）。
    （TextWriter 的颜色在 write_text 时统一指定，因此每页按颜色分别用一个 TextWriter。）"""
    blob = (title or "") + (text or "") + (source_url or "")
    has_cjk = bool(re.search(r'[\u4e00-\u9fff\u3040-\u30ff]', blob))
    body_font = fitz.Font("china-s") if has_cjk else fitz.Font("helv")

    W, H = fitz.paper_size("a4")
    ML = MR = MT = MB = 56.0
    usable = W - ML - MR
    BLACK = (0, 0, 0)
    GRAY = (0.5, 0.5, 0.5)

    def wrap(s, size):
        s = s.replace("\t", "    ")
        out_lines, cur, last_space = [], "", -1
        for ch in s:
            if body_font.text_length(cur + ch, size) <= usable:
                cur += ch
                if ch == " ":
                    last_space = len(cur) - 1
            else:
                if ch == " ":
                    out_lines.append(cur); cur = ""; last_space = -1
                elif last_space >= 0 and (len(cur) - last_space) < 30:
                    out_lines.append(cur[:last_space]); cur = cur[last_space + 1:] + ch; last_space = -1
                else:
                    out_lines.append(cur); cur = ch; last_space = -1
        if cur:
            out_lines.append(cur)
        return out_lines or [""]

    # 第一遍：把所有内容拆成带样式的「行」清单
    items = []  # (text, size, leading, color, gap_after)
    if title:
        for ln in wrap(title, 19):
            items.append([ln, 19, 26, BLACK, 0])
        items[-1][4] = 8.0
    if source_url:
        for ln in wrap(source_url, 8.5):
            items.append([ln, 8.5, 13, GRAY, 0])
        items[-1][4] = 14.0
    for para in (text or "").split("\n"):
        para = para.strip()
        if not para:
            if items:
                items[-1][4] = max(items[-1][4], 8.0)
            continue
        lines = wrap(para, 11.5)
        for i, ln in enumerate(lines):
            items.append([ln, 11.5, 18, BLACK, 6.0 if i == len(lines) - 1 else 0])

    # 第二遍：分页排版。每页对每种颜色各用一个 TextWriter，最后统一 write_text。
    doc = fitz.open()

    def render_page(page_items):
        page = doc.new_page(width=W, height=H)
        writers = {}
        y = MT
        for txt, size, leading, color, gap in page_items:
            tw = writers.get(color)
            if tw is None:
                tw = writers[color] = fitz.TextWriter(page.rect)
            tw.append((ML, y + size), txt, font=body_font, fontsize=size)
            y += leading + gap
        for color, tw in writers.items():
            tw.write_text(page, color=color)

    page_items, y = [], MT
    for it in items:
        line_h = it[2] + it[4]
        if y + it[2] > H - MB and page_items:
            render_page(page_items)
            page_items, y = [], MT
        page_items.append(it)
        y += line_h
    if page_items:
        render_page(page_items)
    if doc.page_count == 0:
        doc.new_page(width=W, height=H)
    doc.save(out_path, deflate=True)
    doc.close()


def build_docx(book, hls, notes, out_path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    doc = Document()
    try:
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass
    doc.add_heading(book["name"], level=0)
    doc.add_paragraph("导出 " + datetime.now().strftime("%Y-%m-%d %H:%M"))
    doc.add_heading("划重点", level=1)
    for h in hls:
        doc.add_paragraph("第 %d 页：%s" % (h["page"], h["text"]))
    doc.add_heading("笔记", level=1)
    for n in notes:
        doc.add_heading(n.get("title", "笔记"), level=2)
        doc.add_paragraph(n.get("body", ""))
    doc.save(out_path)


def build_pdf(book, hls, notes, out_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    ss = getSampleStyleSheet()
    base = ParagraphStyle("cn", parent=ss["Normal"], fontName="STSong-Light", fontSize=11, leading=17)
    h1 = ParagraphStyle("cnh1", parent=base, fontSize=17, leading=23, spaceAfter=10)
    h2 = ParagraphStyle("cnh2", parent=base, fontSize=13, leading=18, spaceBefore=10, spaceAfter=4)
    doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    flow = [Paragraph(_esc(book["name"]), h1),
            Paragraph("导出 " + datetime.now().strftime("%Y-%m-%d %H:%M"), base),
            Spacer(1, 10), Paragraph("划重点", h2)]
    for h in hls:
        flow.append(Paragraph("第 %d 页：%s" % (h["page"], _esc(h["text"])), base))
        flow.append(Spacer(1, 3))
    flow.append(Paragraph("笔记", h2))
    for n in notes:
        flow.append(Paragraph(_esc(n.get("title", "笔记")), h2))
        for line in (n.get("body", "") or "").split("\n"):
            flow.append(Paragraph(_esc(line) or "&nbsp;", base))
    doc.build(flow)


@app.route("/api/export/<bid>")
def api_export(bid):
    b = find_book(bid)
    if not b:
        return "not found", 404
    fmt = request.args.get("fmt", "docx")
    hls = sorted([h for h in load_json(HL_FILE, []) if h["book_id"] == bid], key=lambda x: x["page"])
    notes = notes_list(load_json(NOTE_FILE, {}), bid)
    suffix = ".docx" if fmt == "docx" else ".pdf"
    tmp = DATA_DIR / ("_export" + suffix)
    try:
        if fmt == "docx":
            build_docx(b, hls, notes, str(tmp))
        else:
            build_pdf(b, hls, notes, str(tmp))
    except ImportError:
        need = "python-docx" if fmt == "docx" else "reportlab"
        return jsonify({"error": f"需要先安装 {need}：在 Anaconda Prompt 里运行 pip install {need}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return send_file(str(tmp), as_attachment=True, download_name=(b["name"] + suffix))


# ============================ 前端 ============================
# INDEX_HTML 已外移到 web/index.html、web/style.css、web/app.js（由 / 与 /web/ 路由下发）


# ===========================================================================
# 人文社科 / 跨学科 精读层（内联，不依赖 think_tank；缺它也能运行）
#   1) 个人文献库支持"粘贴观点/文本"条目（不止 PDF）
#   2) 论证框架重建（客观）/ 内在批判（思辨）/ 五维提炼 / 双视野对照
#   3) 文献库定位（陈述观点→定位论据）/ 笔记↔文献库 双向勾连
# 一切引证均要求逐字原文 + 精确出处（条目名 + 第N页/块）。
# ===========================================================================
TEXT_DIR = DATA_DIR / "texts"



def _text_item_path(book):
    p = book.get("path") or ""
    try:
        if p and Path(p).exists():
            return p
    except Exception:
        pass
    return str(TEXT_DIR / (book.get("id", "")))


def _read_text_pages(book, page_chars=1400):
    """把'粘贴文本/片段'类条目读成若干伪页，便于按'第N页/块'标注出处。"""
    raw = ""
    try:
        raw = Path(_text_item_path(book)).read_text(encoding="utf-8")
    except Exception:
        raw = book.get("inline_text", "") or ""
    raw = (raw or "").strip()
    if not raw:
        return [""]
    if book.get("kind") == "snippet":
        return [raw]
    paras = [p.strip() for p in re.split(r"\n\s*\n", raw) if p.strip()]
    pages, cur = [], ""
    for para in paras:
        if cur and len(cur) + len(para) > page_chars:
            pages.append(cur)
            cur = para
        else:
            cur = (cur + "\n\n" + para) if cur else para
    if cur:
        pages.append(cur)
    return pages or [raw]


_PLABEL_CACHE = {}


def _page_labels(book):
    """印刷页码（文本内页码）映射：阅读序号(1..N) -> 标签字符串。
    优先 PDF 内嵌页签(get_label)；否则用每本书的偏移 page_offset；否则回退为阅读序号。"""
    if not book:
        return []
    bid = book.get("id")
    path = book.get("path")
    try:
        mt = os.path.getmtime(path) if path and os.path.exists(path) else 0
    except Exception:
        mt = 0
    off = int(book.get("page_offset") or 0)
    ck = _PLABEL_CACHE.get(bid)
    if ck and ck[0] == (mt, off):
        return ck[1]
    labels = []
    try:
        doc = fitz.open(path)
        pc = doc.page_count
        try:
            has_embedded = bool(doc.get_page_labels())
        except Exception:
            has_embedded = False
        raw = []
        if has_embedded:
            for i in range(pc):
                try:
                    lab = (doc[i].get_label() or "").strip()
                except Exception:
                    lab = ""
                raw.append(lab or str(i + 1))
        doc.close()
        for i in range(pc):
            if has_embedded:
                labels.append(raw[i])
            elif off:
                pn = i + 1 - off
                labels.append(str(pn) if pn >= 1 else str(i + 1))
            else:
                labels.append(str(i + 1))
    except Exception:
        labels = []
    _PLABEL_CACHE[bid] = ((mt, off), labels)
    return labels


def _plab(book, idx1):
    """阅读序号(1-based) -> 印刷页码标签；越界或无映射时回退为该序号。"""
    try:
        labs = _page_labels(book)
        if 1 <= int(idx1) <= len(labs):
            return labs[int(idx1) - 1]
    except Exception:
        pass
    return str(idx1)


def _parse_pages(spec, n):
    """'1-3' / '2,5,7' / '1-3,6' -> 已排序去重、落在 [1,n] 的页码列表。"""
    out = set()
    for part in (spec or "").replace("，", ",").replace("－", "-").split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                a, b = int(a), int(b)
                for x in range(min(a, b), max(a, b) + 1):
                    if 1 <= x <= n:
                        out.add(x)
            except Exception:
                continue
        else:
            try:
                x = int(part)
                if 1 <= x <= n:
                    out.add(x)
            except Exception:
                continue
    return sorted(out)


def _hum_scope_text(book, d, max_chars=48000):
    """按 scope 取文本：all=全文 / page=当前页 / range=页码范围。返回 (文本, 范围标签)。"""
    scope = d.get("scope") or "all"
    pages = book_pages_text(book)
    n = len(pages)
    if scope == "page":
        pno = max(1, min(int(d.get("page") or 1), n))
        return clean_text(pages[pno - 1] or "").strip(), f"第 {pno} 页"
    if scope == "range":
        sel = _parse_pages(d.get("page_range") or "", n)
        if not sel:
            return "", ""
        parts = [f"[第{_plab(book, p)}页]\n{clean_text(pages[p - 1] or '').strip()}" for p in sel]
        return ("\n\n".join(parts))[:max_chars], "第 " + (d.get("page_range") or "").strip() + " 页"
    parts = [f"[第{_plab(book, i + 1)}页/块]\n{clean_text(pages[i] or '').strip()}"
             for i in range(n) if (pages[i] or "").strip()]
    return ("\n\n".join(parts))[:max_chars], "全文"


def _library_corpus(exclude_id=None, max_chars=60000, per_item=8000):
    """把整个个人文献库（含粘贴文本/片段）拼成带出处标记的语料；按预算截断。"""
    out, budget = [], max_chars
    for b in safe_library():
        if exclude_id and b.get("id") == exclude_id:
            continue
        try:
            pages = book_pages_text(b)
        except Exception:
            continue
        parts = [f"[第{_plab(book, i + 1)}页/块] {clean_text(p or '').strip()}"
                 for i, p in enumerate(pages) if (p or "").strip()]
        if not parts:
            continue
        chunk = f"【出处：{b.get('name', '?')}】\n" + "\n".join(parts)[:per_item]
        if budget - len(chunk) < 0:
            out.append(chunk[:max(0, budget)])
            break
        out.append(chunk)
        budget -= len(chunk)
    return "\n\n----\n\n".join(out)


def _notes_corpus(max_chars=40000):
    """把个人笔记库拼成带出处（笔记标题 / 所属文献）的语料。"""
    alln = load_json(NOTE_FILE, {}) or {}
    libmap = {b.get("id"): b.get("name", "?") for b in safe_library()}
    out, budget = [], max_chars
    for bid in alln:
        for nt in notes_list(alln, bid):
            body = (nt.get("body") or "").strip()
            if not body:
                continue
            chunk = f"【笔记：{nt.get('title') or '笔记'}｜出自《{libmap.get(bid, bid)}》】\n{body}"
            if budget - len(chunk) < 0:
                out.append(chunk[:max(0, budget)])
                return "\n\n".join(out)
            out.append(chunk)
            budget -= len(chunk)
    return "\n\n".join(out)


def _horizon_clause(h):
    h = (h or "").strip()
    if h == "en":
        return ("\n\n【视野设定】请在【英语学界】的问题意识与理论框架下解读："
                "只参照英文发表的研究与英语世界的概念坐标，它会被放进哪些英语学界的辩论与对话者之中。")
    if h == "zh":
        return ("\n\n【视野设定】请在【中文学界】的问题意识与理论框架下解读："
                "涵盖一切中文文献（无论简体或繁体，无论大陆、台港或海外华语学界），"
                "它在中文/中国学术语境里介入的是哪些辩论、关切与概念坐标；"
                "注意中文学界与外语学界对'什么算重要问题'的判断可能不同。")
    if h == "de_fr":
        return ("\n\n【视野设定】请在【德语与法语学界（欧陆传统）】的问题意识与理论框架下解读："
                "主要参照德语和法语发表的文献与其概念坐标（如现象学、批判理论/法兰克福、"
                "结构主义与后结构主义、欧陆社会理论与哲学传统），指出它会进入哪些德法学界的辩论与对话者；"
                "凡涉及英语学界视角的投射，请标注'(不确定)'。")
    if h == "foreign":
        return ("\n\n【视野设定】请在【外文（非中文）学界】的问题意识与理论框架下解读："
                "涵盖德语、日语、西班牙语、法语与英语等多语种发表的研究与各自的问题意识，"
                "不要只用英语视角以偏概全；尽量分辨不同语言学术传统看重或忽略的东西，"
                "凡某一语种的把握不确定处，请标注'(不确定)'。")
    return ""


def _hum_need(d):
    """统一取书 + 校验引擎；返回 (book, provider, key, err_response_or_None)。"""
    b = find_book(d.get("book_id")) if d.get("book_id") else None
    if not b:
        return None, None, None, (jsonify({"error": "请先选择一篇文献/文本（在'人文社科精读'窗口的下拉里选）"}), 400)
    provider, key = resolve_key(d)
    if not key:
        return None, None, None, (jsonify({"error": "请先在设置里填入当前引擎的 API Key"}), 400)
    return b, provider, key, None


def _sse(provider, key, sys, user, d):
    return sse_stream(llm_chat_stream(provider, key, sys,
                                      [{"role": "user", "content": user}],
                                      model=d.get("model"), effort=d.get("effort"),
                                      base_url=d.get("base_url")))


# ---- 个人文献库：粘贴导入观点 / 文本（不止 PDF） ----
@app.route("/api/lib/add_text", methods=["POST"])
def api_lib_add_text():
    d = request.get_json(force=True)
    text = (d.get("text") or "").strip()
    name = (d.get("name") or "").strip()
    kind = d.get("kind") if d.get("kind") in ("text", "snippet") else "text"
    if not text:
        return jsonify({"error": "内容为空"}), 400
    if not name:
        name = (text[:24] + ("…" if len(text) > 24 else "")) or "未命名片段"
    try:
        TEXT_DIR.mkdir(parents=True, exist_ok=True)
        tid = ("snip_" if kind == "snippet" else "txt_") + uuid.uuid4().hex[:10] + ".txt"
        dest = TEXT_DIR / tid
        dest.write_text(text, encoding="utf-8")
        book = {
            "id": tid, "name": name, "path": str(dest), "kind": kind,
            "total_pages": len(_read_text_pages({"id": tid, "path": str(dest), "kind": kind})),
            "current_page": 0, "added": datetime.now().strftime("%Y-%m-%d"),
        }
        lib = safe_library()
        lib.append(book)
        save_json(LIB_FILE, lib)
        _text_cache.pop(tid, None)
        return jsonify(book)
    except Exception:
        append_log("/api/lib/add_text 异常", traceback.format_exc())
        return jsonify({"error": "保存失败；详情见 data/logs"}), 500


# ---- 客观分析：研究问题/关键概念/理论方法/论证形式与逻辑/所用材料·史实·案例（只描述不评价） ----
@app.route("/api/hum/framework_stream", methods=["POST"])
def api_hum_framework_stream():
    d = request.get_json(force=True)
    b, provider, key, err = _hum_need(d)
    if err:
        return err
    text, label = _hum_scope_text(b, d)
    if not text:
        return jsonify({"error": "该范围没有可提取文本（扫描件需先 OCR）"}), 400
    sys = ("你是严谨克制的学术分析者。请【客观分析】这一文本——只描述、不评价，不作褒贬、不作批判，"
           "也不替作者补全或拔高；忠于文本本身，凡不确定处标注（不确定）。" + IRON)
    sys = with_glossary(sys, d.get("glossary"))
    user = ("请客观分析以下文本，沿五个维度组织（某一维若文本未涉及，直接跳过、不要硬凑）：\n"
            "① 研究问题：作者真正要回答或处理的问题是什么。\n"
            "② 关键概念的界定与用法：文本中起支撑作用的核心概念，各自如何被界定、在文内如何使用（保留原文关键词）。\n"
            "③ 理论与方法：所依托的理论传统/思想资源，以及处理材料的方法（文本与史料考证、规范分析、案例、比较、统计数据等——有数据才谈数据）。\n"
            "④ 论证形式与逻辑：论证采取的形式（叙述/思辨/形式化/实证），以及它如何一步步展开——前提与预设→关键步骤→结论；尽量标出推断链，并沿用文本自身的展开方式（围绕史料、围绕事件、围绕某项政策的评价或围绕数据）。\n"
            "⑤ 所用材料·史实·案例：作者据以立论的材料、史实与例证。\n"
            f"\n用简体中文，按维度组织，缺失的维度直接不写。\n\n【范围：{label}】\n" + text)
    return _sse(provider, key, sys, user, d)


# ---- 深度阅读：内在批判(根基)＋隐含价值/反身性＋概念史＋谱系学，织成连贯思辨短论；可锁定视角 ----
@app.route("/api/hum/critique_stream", methods=["POST"])
def api_hum_critique_stream():
    d = request.get_json(force=True)
    b, provider, key, err = _hum_need(d)
    if err:
        return err
    text, label = _hum_scope_text(b, d)
    if not text:
        return jsonify({"error": "该范围没有可提取文本（扫描件需先 OCR）"}), 400
    persp = (d.get("perspective") or "").strip()
    sys = ("你是思辨型的批判理论家，以法兰克福学派的【内在批判】为根基审读人文社科文本。"
           "你不逐字复述，而是从哲学论辩与思辨的高度把握文本、再展开深度阅读——"
           "【思辨性与理论张力是你的首要追求】，鼓励辩证、对话式的展开，而非要点罗列。"
           "【请先以最强意义重构作者的论证，再施加批判】。" + IRON)
    task = ("请把以下四条线索织成一篇连贯的辩证短论（不要机械分点罗列，让它们彼此呼应、形成张力）：\n"
            "一、内在批判（根基）：用文本【自身】设定的概念与标准衡量它——其声称的承诺与实际论证之间是否有裂隙？"
            "是否依赖了它所反对的东西？关键概念在文本内部是否自我瓦解？这是整篇的根基。\n"
            "二、隐含的价值判断与反身性：揭示文本未明言却预设的价值立场、规范前提与视角；"
            "并反身追问作者（乃至作为读者的你）站在何处发问、这一位置如何形塑了结论。\n"
            "三、概念史：文本的核心概念从何而来、在历史中含义如何流变，作者取用的是其中哪一层、又遮蔽了哪些层。\n"
            "四、谱系学：这些概念、问题意识与评判标准是在怎样的权力与话语条件下被生产出来的，"
            "其“自然而然”有着怎样的偶然来历——以此撬动文本的自明性。\n"
            "让概念史与谱系学成为批判的杠杆，凸显文本内部、以及它与其他思想传统之间的理论张力与思辨论战。")
    if persp:
        task += (f"\n\n【读者锁定的解读视角】{persp}\n"
                 "请额外并优先从这一视角作思辨性的展开，但不要因此牺牲对文本的忠实。")
    sys = with_glossary(sys, d.get("glossary"))
    user = (task + f"\n\n用简体中文，写成连贯思辨的整篇文字。\n\n【范围：{label}】\n" + text)
    return _sse(provider, key, sys, user, d)




# ---- 视野对照：中文视野 / 外文视野(英·德·法) 两栏对照，对照小结置底 ----
@app.route("/api/hum/dual_horizon_stream", methods=["POST"])
def api_hum_dual_horizon_stream():
    d = request.get_json(force=True)
    b, provider, key, err = _hum_need(d)
    if err:
        return err
    text, label = _hum_scope_text(b, d)
    if not text:
        return jsonify({"error": "该范围没有可提取文本（扫描件需先 OCR）"}), 400
    sys = ("你是跨语言、跨学术传统的解读者，擅长分辨不同学界“问题意识”的差异。" + IRON)
    user = ("请把这一文本放进【两种学术视野】里对照解读，只分两栏，最后给对照小结。"
            "请严格按下面三个小标题输出（先两栏视野，对照小结置于最后）：\n\n"
            "## 中文视野\n在中文学界（涵盖简体与繁体、大陆与台港及海外华语学界）的问题意识与理论坐标下，"
            "这一文本会被放进哪些辩论、对话者与概念坐标。\n\n"
            "## 外文视野\n在英语与欧陆（英·德·法）学界的问题意识与理论框架下——如分析传统、现象学、"
            "批判理论/法兰克福、结构主义与后结构主义、欧陆社会理论与哲学——这一文本会进入哪些辩论与对话者。\n\n"
            "## 对照小结\n明确指出：在哪些点上“视野一换、意义就变”，两种视野各自会看重或忽略什么。"
            "提醒：你对中文问题意识的把握本身可能带有外语视角的投射，凡不确定处标注（不确定）。\n"
            f"\n用简体中文。\n\n【范围：{label}】\n" + text)
    return _sse(provider, key, sys, user, d)




# ---- 笔记 → 当前阅读：找出与当前阅读相关的、我过去写过的笔记，并一起分析 ----
@app.route("/api/links/notes_from_reading_stream", methods=["POST"])
def api_links_notes_from_reading_stream():
    d = request.get_json(force=True)
    b, provider, key, err = _hum_need(d)
    if err:
        return err
    text, label = _hum_scope_text(b, d if d.get("scope") else {"scope": "page", "page": d.get("page", 1)})
    notes = _notes_corpus()
    if not notes:
        return jsonify({"error": "还没有任何笔记"}), 400
    sys = ("你是擅长把'当前阅读'与'读者既有笔记'勾连起来的学术助手。" + IRON)
    user = ("【当前阅读文本（范围：" + label + "）】\n" + text +
            "\n\n【读者的个人笔记库】\n" + notes +
            "\n\n任务：找出与当前阅读相关的、读者过去写过的观点/笔记，并把二者放在一起分析"
            "（呼应、张力、可深化处）。引用笔记与原文都要逐字摘录并标出处（笔记标题 / 第N页）。"
            "无相关笔记就明说。用简体中文。")
    return _sse(provider, key, sys, user, d)


# ---- 笔记 → 文献库：依据我的笔记/观点，在文献库中寻找相关论据 ----
@app.route("/api/links/evidence_from_notes_stream", methods=["POST"])
def api_links_evidence_from_notes_stream():
    d = request.get_json(force=True)
    provider, key = resolve_key(d)
    if not key:
        return jsonify({"error": "请先在设置里填入当前引擎的 API Key"}), 400
    notes = _notes_corpus(max_chars=20000)
    if not notes:
        return jsonify({"error": "还没有任何笔记"}), 400
    corpus = _library_corpus(max_chars=50000)
    if not corpus:
        return jsonify({"error": "文献库为空"}), 400
    sys = ("你是擅长用读者笔记去检索文献库证据的学术助手。" + IRON)
    user = ("【读者的笔记/观点】\n" + notes +
            "\n\n【个人文献库】\n" + corpus +
            "\n\n任务：在文献库中为这些笔记观点寻找相关的论点/论据（支持或反驳）。每条给出：逐字原文 + "
            "精确出处（条目名 + 第N页/块）+ 关系说明；没有就明说。用简体中文。")
    return _sse(provider, key, sys, user, d)





if __name__ == "__main__":
    url = "http://127.0.0.1:5000"
    threading.Timer(1.5, lambda: webbrowser.open(url)).start()
    print("启动中…浏览器将自动打开：", url)
    print("（保持这个窗口开着；想停止就按 Ctrl+C 或关掉窗口）")
    app.run(host="127.0.0.1", port=5000, debug=False)